from datetime import timedelta
from typing import Any

import matplotlib.pyplot as plt
import numpy as np
import psycopg2.extensions
from docx import Document
from docx.shared import Inches

from app.core.exceptions.base import NotFoundException


def generate_report(cursor: psycopg2.extensions.cursor, save_path: str) -> None:
    document = Document()

    document.add_heading("Library Report", 0)

    avg_book_returnment_time = get_average_returnment_time(cursor)

    if avg_book_returnment_time is not None:
        p = document.add_paragraph("В среднем, книги возвращают за ")
        p.add_run(str(avg_book_returnment_time.days)).bold = True
        p.add_run(" д. ")
        p.add_run(str(avg_book_returnment_time.seconds // 3600)).bold = True
        p.add_run(" ч.")

    p = document.add_paragraph("Средний возраст читателя составляет ")
    p.add_run(str(get_average_reader_age(cursor))).bold = True
    p.add_run(" лет.")

    book_to_quantity = get_most_popular_books(cursor)
    plt_save_path = f"{save_path}.jpg"

    index = np.arange(len(book_to_quantity))
    values = list(book_to_quantity.values())

    plt.bar(index, values)
    plt.xticks(index, range(1, len(book_to_quantity) + 1))
    plt.title("Самые читаемые книги")
    plt.xlabel("Номер книги")
    plt.ylabel("Количество читателей")
    plt.savefig(plt_save_path)

    document.add_picture(plt_save_path, Inches(4.5))

    p = document.add_paragraph("Соответствие номера с названием книги: \n")
    for index, book_title in enumerate(book_to_quantity.keys(), start=1):
        p.add_run(f"{index} - {book_title}\n")

    p = document.add_paragraph("Таблица ")
    p.add_run("нарушителей").bold = True
    p.add_run(".")

    column_names = ("ID", "Имя", "Фамилия", "Количество задолженностей")

    table = document.add_table(rows=1, cols=len(column_names))
    hdr_cells = table.rows[0].cells
    for col_index, col_name in enumerate(column_names, start=0):
        hdr_cells[col_index].text = col_name

    for user_id, first_name, last_name, book_count in get_offenders_debts(cursor):
        row_cells = table.add_row().cells

        row_cells[0].text = str(user_id)
        row_cells[1].text = first_name
        row_cells[2].text = last_name
        row_cells[3].text = str(book_count)

    document.save(save_path)


def get_average_returnment_time(cursor: psycopg2.extensions.cursor) -> timedelta | None:
    cursor.execute("""SELECT AVG(age(date_finished, date_start)) FROM book_order WHERE date_finished IS NOT NULL""")
    record = cursor.fetchone()
    if not record:
        raise NotFoundException

    return record[0]


def get_average_reader_age(cursor: psycopg2.extensions.cursor) -> float | None:
    cursor.execute("""SELECT ROUND(AVG(age), 1) FROM user_ WHERE role = 1""")
    record = cursor.fetchone()
    if not record:
        raise NotFoundException

    return record[0]


def get_most_popular_books(cursor: psycopg2.extensions.cursor) -> dict[str, int]:
    sql = """
        SELECT b.title, COUNT(*) AS quantity
        FROM book_order bo
        JOIN book b ON b.id = bo.book_id
        GROUP BY b.title
        ORDER BY quantity DESC
        LIMIT 5
    """
    cursor.execute(sql)
    records = cursor.fetchall()

    return {record["title"]: record["quantity"] for record in records}  # type: ignore


def get_offenders_debts(cursor: psycopg2.extensions.cursor) -> list[tuple[Any, ...]]:
    sql = """
        SELECT u.id, u.first_name, u.last_name, COUNT(*) AS book_count
        FROM offender o
        JOIN user_ u ON o.user_id = u.id
        GROUP BY u.id, u.first_name, u.last_name
    """
    cursor.execute(sql)

    return cursor.fetchall()
